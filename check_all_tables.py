from docx import Document
import os

def check_all_tables(doc_path):
    doc = Document(doc_path)
    
    print(f"\n文件: {os.path.basename(doc_path)}")
    
    # 检查所有段落
    print(f"段落数: {len(doc.paragraphs)}")
    
    # 检查所有表格
    print(f"表格数: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i}:")
        print(f"  行数: {len(table.rows)}")
        print(f"  列数: {len(table.columns)}")
        
        # 打印表格内容
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"  [{row_idx},{col_idx}]: {text[:100]}")
    
    # 检查是否有嵌套表格
    print(f"\n检查嵌套表格:")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"  段落 {i}: {paragraph.text[:80]}")

# 检查一个样本文件
sample_files = [
    '公厕/中医院公厕.docx',
    '垃圾站/西水岸垃圾站.docx',
    '乡镇中转站/龙台镇双龙垃圾压缩中转站.docx'
]

for file_path in sample_files:
    if os.path.exists(file_path):
        check_all_tables(file_path)