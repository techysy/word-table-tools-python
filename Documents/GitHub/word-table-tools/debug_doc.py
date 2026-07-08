# -*- coding: utf-8 -*-
import os
import glob
from docx import Document

def debug_doc(file_path):
    doc = Document(file_path)
    file_name = os.path.basename(file_path)
    print(f"\n{'='*60}")
    print(f"文件: {file_name}")
    print(f"表格数: {len(doc.tables)}")
    print(f"段落数: {len(doc.paragraphs)}")
    
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- 表格 {t_idx}: {len(table.rows)}行 x {len(table.columns)}列 ---")
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    # 显示 repr 以查看隐藏字符
                    print(f"  [{r_idx},{c_idx}] repr: {repr(text[:120])}")
                    # 显示拼接后的文本
                    concat = "".join(p.text for p in cell.paragraphs)
                    print(f"  [{r_idx},{c_idx}] concat: {repr(concat[:120])}")
                    # 检查是否包含"业主"或"承建"
                    if "业主" in text or "承建" in text or "单位" in text:
                        print(f"  *** 包含关键词 ***")
                    print()

# 扫描所有docx文件
base = os.path.dirname(os.path.abspath(__file__))
for root, dirs, files in os.walk(base):
    for f in files:
        if f.endswith('.docx') and not f.startswith('~'):
            debug_doc(os.path.join(root, f))
            break  # 只看第一个文件
    break