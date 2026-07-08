from docx import Document
import xml.etree.ElementTree as ET

doc_path = '公厕/中医院公厕.docx'
doc = Document(doc_path)

# 打印文档的XML
print("文档XML结构:")
print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# 检查文档的body
body = doc.element.body
print(f"\nBody子元素数量: {len(list(body))}")

for i, child in enumerate(body):
    print(f"子元素 {i}: {child.tag}")
    if child.tag.endswith('}p'):  # 段落
        print(f"  段落内容: {child.text}")
    elif child.tag.endswith('}tbl'):  # 表格
        print(f"  表格")
        # 检查表格的行
        rows = child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
        print(f"  表格行数: {len(rows)}")
        for row_idx, row in enumerate(rows):
            cells = row.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
            print(f"    行 {row_idx}: {len(cells)} 个单元格")
            for cell_idx, cell in enumerate(cells):
                # 获取单元格文本
                paragraphs = cell.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
                text = ""
                for p in paragraphs:
                    runs = p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    for run in runs:
                        if run.text:
                            text += run.text
                if text:
                    print(f"      单元格 {cell_idx}: {text[:50]}")