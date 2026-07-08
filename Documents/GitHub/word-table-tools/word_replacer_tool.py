# -*- coding: utf-8 -*-
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import glob
from datetime import datetime
from docx import Document

class WordReplacerTool:
    def __init__(self, root):
        self.root = root
        self.root.title("Word 表格批量替换工具")
        self.root.geometry("750x600")
        self.root.minsize(650, 500)
        
        self.style = ttk.Style("cosmo")
        
        main_frame = ttk.Frame(root, padding=15)
        main_frame.pack(fill=BOTH, expand=True)
        
        # 标题
        title_frame = ttk.Frame(main_frame)
        title_frame.pack(fill=X, pady=(0, 10))
        ttk.Label(title_frame, text="Word 表格批量替换工具", 
                 font=("Microsoft YaHei", 18, "bold")).pack(side=LEFT)
        ttk.Label(title_frame, text="点击下方区域或拖拽文件到此处添加", 
                 font=("Microsoft YaHei", 9), bootstyle="secondary").pack(side=RIGHT, pady=5)
        
        # 替换规则
        rules_lf = ttk.Labelframe(main_frame, text=" 替换规则 ", bootstyle="info")
        rules_lf.pack(fill=X, pady=(0, 10))
        rules_frame = ttk.Frame(rules_lf, padding=12)
        rules_frame.pack(fill=X)
        
        group1 = ttk.Frame(rules_frame)
        group1.pack(fill=X, pady=3)
        ttk.Label(group1, text="查找:", font=("Microsoft YaHei", 10), width=6).pack(side=LEFT)
        self.find1_entry = ttk.Entry(group1, font=("Microsoft YaHei", 10))
        self.find1_entry.pack(side=LEFT, fill=X, padx=(5, 15))
        ttk.Label(group1, text="替换:", font=("Microsoft YaHei", 10), width=6).pack(side=LEFT)
        self.replace1_entry = ttk.Entry(group1, font=("Microsoft YaHei", 10))
        self.replace1_entry.pack(side=LEFT, fill=X, expand=True, padx=(5, 0))
        
        group2 = ttk.Frame(rules_frame)
        group2.pack(fill=X, pady=3)
        ttk.Label(group2, text="查找:", font=("Microsoft YaHei", 10), width=6).pack(side=LEFT)
        self.find2_entry = ttk.Entry(group2, font=("Microsoft YaHei", 10))
        self.find2_entry.pack(side=LEFT, fill=X, padx=(5, 15))
        ttk.Label(group2, text="替换:", font=("Microsoft YaHei", 10), width=6).pack(side=LEFT)
        self.replace2_entry = ttk.Entry(group2, font=("Microsoft YaHei", 10))
        self.replace2_entry.pack(side=LEFT, fill=X, expand=True, padx=(5, 0))
        
        ttk.Label(rules_frame, text="包含匹配：单元格内容包含查找文字即替换（自动排除表头等长文本），如「业主单位：」可匹配「业主单位：中国移动」",
                 font=("Microsoft YaHei", 8), bootstyle="secondary").pack(anchor=W, pady=(5,0))
        
        # 文件列表
        file_lf = ttk.Labelframe(main_frame, text=" 文件列表 ", bootstyle="success")
        file_lf.pack(fill=BOTH, expand=True, pady=(0, 10))
        file_frame = ttk.Frame(file_lf, padding=10)
        file_frame.pack(fill=BOTH, expand=True)
        
        # 拖拽区域
        self.drop_frame = tk.Frame(file_frame, bg="#e8f5e9", relief="groove", bd=2, cursor="hand2")
        self.drop_frame.pack(fill=BOTH, expand=True)
        
        self.drop_label = tk.Label(self.drop_frame, 
                                   text="拖拽Word文件到此处\n或点击此处添加文件",
                                   bg="#e8f5e9", fg="#2e7d32",
                                   font=("Microsoft YaHei", 12))
        self.drop_label.pack(expand=True)
        
        self.drop_frame.bind("<Button-1>", lambda e: self.add_files())
        self.drop_label.bind("<Button-1>", lambda e: self.add_files())
        
        # 尝试启用拖拽
        self._setup_dnd()
        
        # 文件列表（隐藏，用于内部管理）
        self.file_listbox = tk.Listbox(file_frame, height=6, selectmode=EXTENDED,
                                       font=("Consolas", 9))
        self.file_listbox.pack(fill=BOTH, expand=True, pady=(5, 0))
        
        btn_frame = ttk.Frame(file_frame)
        btn_frame.pack(fill=X, pady=(5, 0))
        ttk.Button(btn_frame, text="添加文件", command=self.add_files, 
                  bootstyle="outline", width=10).pack(side=LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text="添加文件夹", command=self.add_folder,
                  bootstyle="outline", width=10).pack(side=LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text="清空列表", command=self.clear_list,
                  bootstyle="danger-outline", width=10).pack(side=LEFT)
        
        self.count_var = tk.StringVar(value="共 0 个文件")
        ttk.Label(btn_frame, textvariable=self.count_var, 
                 font=("Microsoft YaHei", 9), bootstyle="secondary").pack(side=RIGHT)
        
        # 底部按钮
        bottom_frame = ttk.Frame(main_frame)
        bottom_frame.pack(fill=X)
        ttk.Button(bottom_frame, text="执行替换", command=self.execute_replace,
                  bootstyle="success", width=15).pack(side=LEFT, padx=(0, 10))
        ttk.Button(bottom_frame, text="预览替换", command=self.preview_replace,
                  bootstyle="info-outline", width=15).pack(side=LEFT)
        
        self.status_var = tk.StringVar(value="就绪")
        ttk.Label(bottom_frame, textvariable=self.status_var, 
                 font=("Microsoft YaHei", 9), bootstyle="secondary").pack(side=RIGHT)
        
        self.files = []
        self.log_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logs")
        os.makedirs(self.log_dir, exist_ok=True)
        
    def _setup_dnd(self):
        """设置拖拽功能"""
        try:
            from tkinterdnd2 import DND_FILES
            self.root.drop_target_register(DND_FILES)
            self.root.dnd_bind('<<Drop>>', self._on_drop)
            self.drop_label.config(text="拖拽Word文件到此处\n或点击此处添加文件")
        except Exception as e:
            self.drop_label.config(text="点击此处添加Word文件")
            
    def _on_drop(self, event):
        """处理拖拽文件"""
        try:
            data = event.data
            # 处理可能的多文件格式
            files = []
            if '{' in data:
                # 格式: {file1} {file2}
                import re
                files = re.findall(r'\{([^}]+)\}', data)
                if not files:
                    files = [data.strip('{}')]
            else:
                files = data.split()
            
            for file_path in files:
                file_path = file_path.strip()
                if file_path.lower().endswith('.docx') and file_path not in self.files:
                    self.files.append(file_path)
                    self.file_listbox.insert(END, os.path.basename(file_path))
            self.count_var.set(f"共 {len(self.files)} 个文件")
            self._update_drop_label()
        except Exception as e:
            print(f"拖拽处理错误: {e}")
        
    def _update_drop_label(self):
        if self.files:
            self.drop_label.config(text=f"已添加 {len(self.files)} 个文件\n继续拖拽或点击添加更多")
        else:
            self.drop_label.config(text="拖拽Word文件到此处\n或点击此处添加文件")
        
    def log(self, msg):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        line = f"[{timestamp}] {msg}"
        print(line)
        log_file = os.path.join(self.log_dir, f"replace_{datetime.now().strftime('%Y%m%d')}.log")
        with open(log_file, "a", encoding="utf-8") as f:
            f.write(line + "\n")
            
    def log_cell(self, file_name, table_idx, row_idx, col_idx, cell_text, action, new_text=""):
        detail = f"文件:{file_name} | 表格{table_idx} 行{row_idx}列{col_idx} | 内容:[{cell_text[:80]}] | {action}"
        if new_text:
            detail += f" -> [{new_text[:80]}]"
        self.log(detail)
        
    def add_files(self):
        filetypes = [("Word文件", "*.docx"), ("所有文件", "*.*")]
        files = filedialog.askopenfilenames(filetypes=filetypes)
        for file in files:
            if file not in self.files:
                self.files.append(file)
                self.file_listbox.insert(END, os.path.basename(file))
        self.count_var.set(f"共 {len(self.files)} 个文件")
        self._update_drop_label()
        
    def add_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            docx_files = glob.glob(os.path.join(folder, "**/*.docx"), recursive=True)
            for file in docx_files:
                if file not in self.files:
                    self.files.append(file)
                    self.file_listbox.insert(END, os.path.basename(file))
            self.count_var.set(f"共 {len(self.files)} 个文件")
            self._update_drop_label()
            
    def clear_list(self):
        self.files.clear()
        self.file_listbox.delete(0, END)
        self.count_var.set("共 0 个文件")
        self._update_drop_label()
        
    def replace_cell_content(self, cell, replacements, file_name, table_idx, row_idx, col_idx):
        full_text = ""
        for paragraph in cell.paragraphs:
            full_text += paragraph.text
        full_text = full_text.strip()
        
        if not full_text:
            self.log_cell(file_name, table_idx, row_idx, col_idx, "(空)", "跳过-空单元格")
            return False
        
        for find_text, replace_text in replacements.items():
            find_clean = find_text.strip()
            if find_clean in full_text:
                self.log_cell(file_name, table_idx, row_idx, col_idx, full_text, "包含匹配成功", replace_text)
                for i, paragraph in enumerate(cell.paragraphs):
                    if i == 0:
                        if paragraph.runs:
                            first_run = paragraph.runs[0]
                            for run in paragraph.runs[1:]:
                                run.text = ""
                            first_run.text = replace_text
                        else:
                            paragraph.text = replace_text
                    else:
                        for run in paragraph.runs:
                            run.text = ""
                return True
        
        self.log_cell(file_name, table_idx, row_idx, col_idx, full_text, "无匹配规则")
        return False
        
    def _get_cell_text(self, tc_element):
        """从XML单元格元素获取文本"""
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        paragraphs = tc_element.findall(f'.//{ns}p')
        texts = []
        for p in paragraphs:
            ts = p.findall(f'.//{ns}t')
            text = ''.join([t.text for t in ts if t.text])
            texts.append(text)
        return '\n'.join(texts).strip()
        
    def process_file(self, file_path, replacements):
        """处理单个文件 - 只替换匹配文字，不删除其他内容"""
        import zipfile
        import tempfile
        import shutil
        from lxml import etree
        
        file_name = os.path.basename(file_path)
        self.log(f"开始处理: {file_name}")
        
        try:
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            changes = 0
            
            with zipfile.ZipFile(file_path, 'r') as zin:
                xml_content = zin.read('word/document.xml')
            
            root = etree.fromstring(xml_content)
            
            # 遍历所有表格单元格
            for t_idx, tbl in enumerate(root.findall(f'.//{ns}tbl')):
                for r_idx, tr in enumerate(tbl.findall(f'.//{ns}tr')):
                    for c_idx, tc in enumerate(tr.findall(f'{ns}tc')):
                        # 获取单元格完整文本
                        paragraphs = tc.findall(f'.//{ns}p')
                        cell_text = ""
                        for p in paragraphs:
                            for t in p.findall(f'.//{ns}t'):
                                if t.text:
                                    cell_text += t.text
                        
                        if not cell_text.strip():
                            continue
                        
                        # 检查匹配 - 包含匹配，但排除长文本（表头）
                        for find_text, replace_text in replacements.items():
                            if find_text.strip() in cell_text and len(cell_text.strip()) < 50:
                                self.log_cell(file_name, t_idx, r_idx, c_idx, cell_text.strip(), "匹配成功", replace_text)
                                # 只替换第一个段落的第一个t，其他不动
                                if paragraphs:
                                    first_p = paragraphs[0]
                                    t_list = first_p.findall(f'.//{ns}t')
                                    if t_list:
                                        # 把新文本放到第一个t
                                        t_list[0].text = replace_text
                                        # 清空同一段落其他t（避免重复显示）
                                        for t in t_list[1:]:
                                            t.text = ""
                                changes += 1
                                break
            
            # 处理段落 - 包含匹配，排除长文本
            for p in root.findall(f'.//{ns}body/{ns}p'):
                t_list = p.findall(f'.//{ns}t')
                full_text = ''.join([t.text for t in t_list if t.text]).strip()
                if not full_text:
                    continue
                for find_text, replace_text in replacements.items():
                    if find_text.strip() in full_text and len(full_text) < 50:
                        self.log(f"  段落匹配: [{full_text[:60]}]")
                        t_list[0].text = replace_text
                        for t in t_list[1:]:
                            t.text = ""
                        changes += 1
                        break
            
            if changes > 0:
                new_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
                temp_file = file_path + '.tmp'
                with zipfile.ZipFile(file_path, 'r') as zin:
                    with zipfile.ZipFile(temp_file, 'w', zipfile.ZIP_DEFLATED) as zout:
                        for item in zin.infolist():
                            if item.filename == 'word/document.xml':
                                zout.writestr(item, new_xml)
                            else:
                                zout.writestr(item, zin.read(item.filename))
                shutil.move(temp_file, file_path)
                self.log(f"  保存成功，替换了 {changes} 处")
            else:
                self.log(f"  无匹配内容")
            return True, changes
            
        except Exception as e:
            import traceback
            self.log(f"  处理出错: {e}")
            self.log(traceback.format_exc())
            return False, str(e)
            
    def get_replacements(self):
        replacements = {}
        find1 = self.find1_entry.get().strip()
        replace1 = self.replace1_entry.get().strip()
        if find1:
            replacements[find1] = replace1
        find2 = self.find2_entry.get().strip()
        replace2 = self.replace2_entry.get().strip()
        if find2:
            replacements[find2] = replace2
        return replacements
        
    def execute_replace(self):
        if not self.files:
            messagebox.showwarning("提示", "请先添加文件")
            return
        replacements = self.get_replacements()
        if not replacements:
            messagebox.showwarning("提示", "请填写替换规则")
            return
            
        msg = f"对 {len(self.files)} 个文件执行替换：\n\n"
        for f, r in replacements.items():
            msg += f"「{f}」->「{r}」\n"
        if not messagebox.askyesno("确认替换", msg):
            return
        
        self.log("=" * 60)
        self.log("开始批量替换")
        self.log(f"替换规则: {replacements}")
        self.log(f"文件数量: {len(self.files)}")
            
        success = 0
        fail = 0
        no_change = 0
        
        for i, file_path in enumerate(self.files):
            self.status_var.set(f"处理中 {i+1}/{len(self.files)}")
            self.root.update()
            result, changes = self.process_file(file_path, replacements)
            if result is True and changes > 0:
                success += 1
            elif changes == 0:
                no_change += 1
            else:
                fail += 1
                
        summary = f"替换完成 | 成功:{success} 无匹配:{no_change} 失败:{fail}"
        self.log(summary)
        self.log("=" * 60)
        self.status_var.set("完成")
        messagebox.showinfo("完成", f"{summary}\n\n日志: logs/replace_{datetime.now().strftime('%Y%m%d')}.log")
        
    def preview_replace(self):
        if not self.files:
            messagebox.showwarning("提示", "请先添加文件")
            return
        replacements = self.get_replacements()
        if not replacements:
            messagebox.showwarning("提示", "请填写替换规则")
            return
            
        win = tk.Toplevel(self.root)
        win.title("预览替换结果")
        win.geometry("750x550")
        
        notebook = ttk.Notebook(win)
        notebook.pack(fill=BOTH, expand=True, padx=10, pady=10)
        
        preview_frame = ttk.Frame(notebook)
        notebook.add(preview_frame, text="预览")
        
        text = tk.Text(preview_frame, wrap=WORD, font=("Microsoft YaHei", 10), padx=10, pady=10)
        text.pack(fill=BOTH, expand=True)
        
        text.insert(END, "替换规则（包含匹配）\n")
        text.insert(END, "-" * 50 + "\n\n")
        for f, r in replacements.items():
            text.insert(END, f"查找: {f}\n")
            text.insert(END, f"替换: {r}\n\n")
        text.insert(END, "-" * 50 + "\n\n")
        
        try:
            doc = Document(self.files[0])
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            text.insert(END, f"文件: {os.path.basename(self.files[0])}\n\n")
            
            # 用XML底层解析表格
            body = doc.element.body
            tables = body.findall(f'.//{ns}tbl')
            
            for t_idx, table_elem in enumerate(tables):
                rows = table_elem.findall(f'.//{ns}tr')
                text.insert(END, f"--- 表格 {t_idx+1} ({len(rows)}行) ---\n")
                
                for r_idx, row_elem in enumerate(rows):
                    cells = row_elem.findall(f'{ns}tc')
                    for c_idx, cell_elem in enumerate(cells):
                        cell_text = self._get_cell_text(cell_elem)
                        if not cell_text:
                            continue
                        # 显示内容
                        text.insert(END, f"  行{r_idx}列{c_idx}: {cell_text[:80]}\n", "dim")
                        text.insert(END, f"    repr: {repr(cell_text[:100])}\n", "dim")
                        # 检查匹配 - 包含匹配，排除长文本
                        matched = False
                        for f, r in replacements.items():
                            if f.strip() in cell_text and len(cell_text.strip()) < 50:
                                text.insert(END, f"  [匹配] -> {r}\n", "match")
                                matched = True
                                break
                        if not matched:
                            # 检查是否因为文本太长被排除（表头）
                            for f, r in replacements.items():
                                f_clean = f.strip()
                                if f_clean in cell_text and len(cell_text.strip()) >= 50:
                                    # 部分字符匹配，找差异点
                                    for ci in range(min(len(f_clean), len(cell_text))):
                                        if cell_text[ci] != f_clean[ci]:
                                            text.insert(END, f"  [差异位置{ci}] 查找:'{f_clean[ci]}'(U+{ord(f_clean[ci]):04X}) vs 实际:'{cell_text[ci]}'(U+{ord(cell_text[ci]):04X})\n", "dim")
                                            break
                                    break
                        text.insert(END, "\n")
                text.insert(END, "\n")
            text.tag_config("match", foreground="#dc3545", font=("Microsoft YaHei", 10, "bold"))
            text.tag_config("dim", foreground="#888888", font=("Consolas", 8))
        except Exception as e:
            text.insert(END, f"读取错误: {e}")
        text.config(state=DISABLED)
        
        log_frame = ttk.Frame(notebook)
        notebook.add(log_frame, text="日志")
        
        log_text = tk.Text(log_frame, wrap=WORD, font=("Consolas", 9), padx=10, pady=10, bg="#1e1e1e", fg="#d4d4d4")
        log_text.pack(fill=BOTH, expand=True)
        
        log_file = os.path.join(self.log_dir, f"replace_{datetime.now().strftime('%Y%m%d')}.log")
        if os.path.exists(log_file):
            with open(log_file, "r", encoding="utf-8") as f:
                log_text.insert(END, f.read())
        else:
            log_text.insert(END, "暂无日志记录\n")
        log_text.config(state=DISABLED)

def main():
    root = ttk.Window(title="Word 表格批量替换工具", themename="cosmo", size=(750, 600))
    app = WordReplacerTool(root)
    root.mainloop()

if __name__ == "__main__":
    main()