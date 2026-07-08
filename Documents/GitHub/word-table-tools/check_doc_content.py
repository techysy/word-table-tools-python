from docx import Document
import os

def check_doc_content(doc_path):
    doc = Document(doc_path)
    
    print(f"\n文件: {os.path.basename(doc_path)}")
    print("=" * 50)
    
    # 打印所有段落
    print("段落内容:")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"  段落 {i}: {paragraph.text[:100]}")
    
    # 打印所有表格
    print(f"\n表格内容:")
    for i, table in enumerate(doc.tables):
        print(f"  表格 {i}: {len(table.rows)}行 x {len(table.columns)}列")
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"    [{row_idx},{col_idx}]: {text[:100]}")

# 检查一个样本文件
sample_file = '公厕/中医院公厕.docx'
if os.path.exists(sample_file):
    check_doc_content(sample_file)