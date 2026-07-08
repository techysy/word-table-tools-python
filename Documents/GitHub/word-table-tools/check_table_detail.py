from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def check_table_detail(doc_path):
    doc = Document(doc_path)
    
    print(f"\n文件: {doc_path}")
    print(f"表格数量: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i}:")
        print(f"  行数: {len(table.rows)}")
        print(f"  列数: {len(table.columns)}")
        
        # 检查表格的XML结构
        tbl = table._tbl
        print(f"  表格XML标签: {tbl.tag}")
        
        # 检查表格宽度
        tbl_pr = tbl.tblPr
        if tbl_pr is not None:
            tbl_w = tbl_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
            if tbl_w is not None:
                print(f"  表格宽度类型: {tbl_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')}")
                print(f"  表格宽度值: {tbl_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')}")
        
        # 检查列宽
        for col_idx, col in enumerate(table.columns):
            print(f"  列 {col_idx}:")
            for row_idx, cell in enumerate(col.cells):
                # 检查单元格宽度
                tc = cell._tc
                tc_pr = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
                if tc_pr is not None:
                    tc_w = tc_pr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
                    if tc_w is not None:
                        print(f"    单元格[{row_idx}]宽度: {tc_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')} {tc_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')}")
                
                # 打印单元格内容
                text = cell.text.strip()
                if text:
                    print(f"    内容: {text[:80]}...")

# 检查一个样本文件
sample_file = '公厕/中医院公厕.docx'
if os.path.exists(sample_file):
    check_table_detail(sample_file)