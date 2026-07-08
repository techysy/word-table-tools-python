from docx import Document
import os

doc_path = '公厕/中医院公厕.docx'
doc = Document(doc_path)

print(f"文档: {doc_path}")
print(f"段落数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

for i, table in enumerate(doc.tables):
    print(f"\n表格 {i}:")
    print(f"  行数: {len(table.rows)}")
    print(f"  列数: {len(table.columns)}")
    
    # 打印表格内容
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  [{row_idx},{col_idx}]: {text[:50]}...")
    
    # 检查表格宽度
    if table.rows:
        first_row = table.rows[0]
        if first_row.cells:
            first_cell = first_row.cells[0]
            print(f"  第一个单元格宽度: {first_cell.width}")