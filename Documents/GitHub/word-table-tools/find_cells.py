from docx import Document
import os

def find_cells(doc_path, search_text):
    """
    在文档表格中查找包含指定文本的单元格
    """
    doc = Document(doc_path)
    
    found_cells = []
    
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if search_text in cell.text:
                    found_cells.append({
                        'table': table_idx,
                        'row': row_idx,
                        'col': col_idx,
                        'text': cell.text,
                        'cell': cell
                    })
    
    return found_cells

# 检查一个样本文件
sample_file = '公厕/中医院公厕.docx'
if os.path.exists(sample_file):
    doc = Document(sample_file)
    
    print("查找'业主单位'相关单元格:")
    cells = find_cells(sample_file, '业主单位')
    for cell_info in cells:
        print(f"  表格{cell_info['table']}, 行{cell_info['row']}, 列{cell_info['col']}: {cell_info['text'][:50]}")
    
    print("\n查找'承建单位'相关单元格:")
    cells = find_cells(sample_file, '承建单位')
    for cell_info in cells:
        print(f"  表格{cell_info['table']}, 行{cell_info['row']}, 列{cell_info['col']}: {cell_info['text'][:50]}")